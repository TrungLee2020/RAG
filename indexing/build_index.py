import os
import argparse
import logging
from pathlib import Path
import pandas as pd
import time
import uuid
from typing import List, Dict, Any
from dotenv import load_dotenv

from text_splitter import DocumentChunk, custom_split
from vector_db import FAISSVectorStore

# Thiết lập logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# Tải environment variables
load_dotenv()

# Đường dẫn mặc định
DEFAULT_INDEX_PATH = os.getenv("INDEX_PATH", "../vector_db/hnsw.index")
DEFAULT_DOCUMENTS_PATH = os.getenv("DOCUMENTS_PATH", "../vector_db/documents.pkl")
DEFAULT_MODEL_NAME = os.getenv("EMBEDDING_MODEL", "Alibaba-NLP/gte-multilingual-base")
TOKENIZE_MODEL = os.getenv("TOKENIZE_MODEL", "vinai/phobert-base-v2")


def get_files_from_directory(directory: str, extensions: List[str] = None) -> List[str]:
    """
    Lấy danh sách đường dẫn đến các file trong thư mục và các thư mục con

    Args:
        directory: Thư mục gốc
        extensions: Danh sách các phần mở rộng cần lọc (ví dụ: ['.txt', '.md'])

    Returns:
        List[str]: Danh sách đường dẫn đến các file
    """
    if extensions is None:
        extensions = ['.txt', '.md', '.pdf', '.docx', '.pptx', '.xlsx', '.csv']

    # Chuyển đổi extensions thành lowercase
    extensions = [ext.lower() for ext in extensions]

    files = []
    root_path = Path(directory)

    # Duyệt qua tất cả files trong thư mục và các thư mục con
    for path in root_path.glob('**/*'):
        if path.is_file() and path.suffix.lower() in extensions:
            files.append(str(path))

    # Log thông tin về cấu trúc thư mục đã quét
    dir_structure = {}
    for file in files:
        parent = os.path.dirname(os.path.relpath(file, directory))
        if parent not in dir_structure:
            dir_structure[parent] = 0
        dir_structure[parent] += 1

    logger.info("Cấu trúc thư mục đã quét:")
    for dir_path, count in dir_structure.items():
        logger.info(f"  {dir_path}: {count} files")

    return files


def read_file_content(file_path: str) -> str:
    """
    Đọc nội dung file dựa vào phần mở rộng

    Args:
        file_path: Đường dẫn đến file

    Returns:
        str: Nội dung file
    """
    suffix = Path(file_path).suffix.lower()

    try:
        # Xử lý file dựa trên định dạng
        if suffix == '.pdf':
            # Xử lý PDF
            from process_document.pdf_parse.extractor_pdf import extract_pdf_markdown
            return extract_pdf_markdown(file_path)

        elif suffix == '.docx':
            # Xử lý DOCX
            import docx
            doc = docx.Document(file_path)
            return '\n\n'.join([para.text for para in doc.paragraphs])

        elif suffix in ['.xlsx', '.xls']:
            # Xử lý Excel
            import pandas as pd
            df = pd.read_excel(file_path)
            return df.to_csv(index=False)

        elif suffix == '.csv':
            # Đọc CSV như text
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()

        else:
            # Các file text thông thường
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()

    except Exception as e:
        logger.error(f"Lỗi khi đọc file {file_path}: {e}")
        return f"Lỗi khi đọc file: {str(e)}"


def extract_metadata_from_file(file_path: str) -> Dict[str, Any]:
    """
    Trích xuất metadata từ file

    Args:
        file_path: Đường dẫn đến file

    Returns:
        Dict: Metadata của file
    """
    path = Path(file_path)

    # Tạo document_id độc đáo hơn bằng cách kết hợp thông tin thư mục và tên file
    folder_path = str(path.parent)
    unique_id = f"{folder_path}_{path.stem}"
    document_id = f"doc_{hash(unique_id) % 100000}"  # Tăng từ 10000 lên 100000 để tăng tính độc đáo

    # Metadata cơ bản
    metadata = {
        "source": file_path,
        "filename": path.name,
        "title": path.stem,
        "file_type": path.suffix.lower()[1:],  # Loại bỏ dấu .
        "date_processed": time.strftime("%Y-%m-%d %H:%M:%S"),
        "document_id": document_id,
        "directory": folder_path  # Thêm thông tin về thư mục cha
    }

    # Trích xuất metadata chuyên biệt dựa vào loại file
    suffix = path.suffix.lower()

    try:
        if suffix == '.pdf':
            # Metadata từ PDF
            import fitz  # PyMuPDF
            doc = fitz.open(file_path)
            pdf_metadata = doc.metadata

            if pdf_metadata:
                if pdf_metadata.get('title'):
                    metadata['title'] = pdf_metadata.get('title')
                if pdf_metadata.get('author'):
                    metadata['author'] = pdf_metadata.get('author')
                if pdf_metadata.get('subject'):
                    metadata['subject'] = pdf_metadata.get('subject')
                if pdf_metadata.get('keywords'):
                    metadata['keywords'] = pdf_metadata.get('keywords')

            doc.close()

        elif suffix == '.docx':
            # Metadata từ DOCX
            import docx
            doc = docx.Document(file_path)

            if hasattr(doc, 'core_properties'):
                props = doc.core_properties
                if props.title:
                    metadata['title'] = props.title
                if props.author:
                    metadata['author'] = props.author
                if props.subject:
                    metadata['subject'] = props.subject
                if props.keywords:
                    metadata['keywords'] = props.keywords

    except Exception as e:
        logger.warning(f"Không thể trích xuất metadata từ {file_path}: {e}")

    return metadata


def load_metadata_csv(metadata_path: str) -> pd.DataFrame:
    """
    Tải metadata từ file CSV

    Args:
        metadata_path: Đường dẫn đến file metadata CSV

    Returns:
        pd.DataFrame: DataFrame chứa metadata
    """
    try:
        return pd.read_csv(metadata_path)
    except Exception as e:
        logger.warning(f"Không thể tải metadata từ {metadata_path}: {e}")
        # Tạo DataFrame trống với các cột cơ bản
        return pd.DataFrame(columns=['filepath', 'document_id', 'title', 'author', 'category', 'date', 'source'])


def split_document_into_chunks(text: str, chunk_size: int = 1024, chunk_overlap: int = 256,
                               document_id: str = None, metadata: Dict[str, Any] = None) -> List[DocumentChunk]:
    """
    Chia văn bản thành các đoạn nhỏ

    Args:
        text: Văn bản cần chia
        chunk_size: Kích thước mỗi đoạn (số tokens)
        chunk_overlap: Số tokens chồng lấp giữa các đoạn
        document_id: ID của tài liệu gốc
        metadata: Metadata của tài liệu

    Returns:
        List[DocumentChunk]: Danh sách các đoạn đã chia
    """
    from transformers import AutoTokenizer

    # Tạo document_id nếu không có
    if document_id is None:
        document_id = str(uuid.uuid4())

    # Khởi tạo metadata nếu không có
    metadata = metadata or {}

    # Khởi tạo tokenizer
    tokenizer = AutoTokenizer.from_pretrained(TOKENIZE_MODEL)

    # Tách văn bản bằng custom_split
    text_chunks = custom_split(text, tokenizer)

    # Tạo DocumentChunk cho mỗi đoạn
    chunks = []
    for i, chunk_text in enumerate(text_chunks):
        chunk = DocumentChunk(
            text=chunk_text,
            document_id=document_id,
            chunk_id=f"{document_id}_{i}",
            metadata={
                **metadata,
                "chunk_index": i,
                "total_chunks": len(text_chunks)
            }
        )
        chunks.append(chunk)

    return chunks


def create_vector_database(
        directory: str,
        metadata_path: str = None,
        index_path: str = DEFAULT_INDEX_PATH,
        documents_path: str = DEFAULT_DOCUMENTS_PATH,
        model_name: str = DEFAULT_MODEL_NAME,
        chunk_size: int = 1024,
        chunk_overlap: int = 256,
        extensions: List[str] = None,
        index_type: str = "hnsw"
):
    """
    Tạo vector database từ các file trong thư mục

    Args:
        directory: Thư mục chứa tài liệu
        metadata_path: Đường dẫn đến file CSV chứa metadata (tùy chọn)
        index_path: Đường dẫn lưu FAISS index
        documents_path: Đường dẫn lưu documents
        model_name: Tên model SentenceTransformer
        chunk_size: Kích thước mỗi chunk (số tokens)
        chunk_overlap: Số tokens chồng lấp giữa các chunks
        extensions: Danh sách các phần mở rộng cần lọc
        index_type: Loại FAISS index (hnsw, flat, ivf)
    """

    # Đảm bảo thư mục lưu trữ tồn tại
    os.makedirs(os.path.dirname(index_path), exist_ok=True)

    # Tạo vector store mới
    logger.info(f"Tạo vector store mới với model {model_name} và loại index {index_type}")
    vector_store = FAISSVectorStore(
        model_name=model_name,
        index_type=index_type,
        index_path=index_path if os.path.exists(index_path) else None,
        documents_path=documents_path if os.path.exists(documents_path) else None
    )

    # Lấy danh sách file
    logger.info(f"Quét files trong thư mục: {directory}")
    files = get_files_from_directory(directory, extensions)
    logger.info(f"Tìm thấy {len(files)} files")

    # Tải metadata từ CSV nếu có
    metadata_df = None
    if metadata_path and os.path.exists(metadata_path):
        logger.info(f"Tải metadata từ: {metadata_path}")
        metadata_df = load_metadata_csv(metadata_path)

    # Xử lý từng file
    for i, file_path in enumerate(files):
        try:
            logger.info(f"Xử lý file {i + 1}/{len(files)}: {file_path}")

            # Đọc nội dung file
            content = read_file_content(file_path)

            if not content or len(content.strip()) < 10:
                logger.warning(f"File {file_path} không có nội dung hoặc quá ngắn, bỏ qua")
                continue

            # Trích xuất metadata
            metadata = extract_metadata_from_file(file_path)

            # Cập nhật metadata từ CSV nếu có
            if metadata_df is not None:
                # Thử nhiều định dạng đường dẫn để tìm kiếm khớp trong metadata
                abs_path = os.path.abspath(file_path)
                rel_path = os.path.relpath(file_path, directory)
                filename = os.path.basename(file_path)

                # Thử khớp với nhiều định dạng đường dẫn khác nhau
                file_meta = metadata_df[
                    (metadata_df['filepath'] == rel_path) |
                    (metadata_df['filepath'] == abs_path) |
                    (metadata_df['filepath'] == file_path) |
                    (metadata_df['filepath'] == filename)
                    ]

                if not file_meta.empty:
                    row = file_meta.iloc[0]
                    # Cập nhật các trường từ CSV
                    for col in file_meta.columns:
                        if col != 'filepath' and not pd.isna(row[col]):
                            metadata[col] = row[col]
                else:
                    logger.warning(f"Không tìm thấy metadata cho file: {file_path}")
                    logger.debug(f"Đã thử các đường dẫn: {rel_path}, {filename}, {abs_path}")

            # Chia tài liệu thành chunks
            chunks = split_document_into_chunks(
                text=content,
                chunk_size=chunk_size,
                chunk_overlap=chunk_overlap,
                document_id=metadata.get('document_id'),
                metadata=metadata
            )

            logger.info(f"Chia {file_path} thành {len(chunks)} chunks")

            # Thêm chunks vào vector store
            vector_store.add_documents(chunks)

        except Exception as e:
            logger.error(f"Lỗi khi xử lý file {file_path}: {e}")

    # Lưu vector store
    logger.info(f"Lưu vector store vào: {index_path} và {documents_path}")
    vector_store.save(index_path=index_path, documents_path=documents_path)

    logger.info(f"Đã hoàn thành, tổng số documents: {vector_store.index.ntotal}")


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Tạo vector database từ thư mục chứa tài liệu")

    parser.add_argument("--directory", "-d", type=str, required=True,
                        help="Thư mục chứa tài liệu cần xử lý")

    parser.add_argument("--metadata", "-m", type=str,
                        help="Đường dẫn đến file CSV chứa metadata (tùy chọn)")

    parser.add_argument("--index-path", type=str, default=DEFAULT_INDEX_PATH,
                        help=f"Đường dẫn lưu FAISS index (mặc định: {DEFAULT_INDEX_PATH})")

    parser.add_argument("--documents-path", type=str, default=DEFAULT_DOCUMENTS_PATH,
                        help=f"Đường dẫn lưu documents (mặc định: {DEFAULT_DOCUMENTS_PATH})")

    parser.add_argument("--model-name", type=str, default=DEFAULT_MODEL_NAME,
                        help=f"Tên model SentenceTransformer (mặc định: {DEFAULT_MODEL_NAME})")

    parser.add_argument("--chunk-size", type=int, default=1024,
                        help="Kích thước mỗi chunk (số tokens) (mặc định: 1024)")

    parser.add_argument("--chunk-overlap", type=int, default=256,
                        help="Số tokens chồng lấp giữa các chunks (mặc định: 256)")

    parser.add_argument("--extensions", type=str, default=".txt,.md,.pdf,.docx",
                        help="Danh sách các phần mở rộng file cần xử lý, phân tách bằng dấu phẩy (mặc định: .txt,.md,.pdf,.docx)")

    parser.add_argument("--index-type", type=str, default="hnsw", choices=["flat", "hnsw", "ivf"],
                        help="Loại FAISS index (mặc định: hnsw)")

    args = parser.parse_args()
    import sys

    # Kiểm tra thư mục đầu vào
    if not os.path.exists(args.directory):
        logger.error(f"Thư mục {args.directory} không tồn tại!")
        sys.exit(1)

    # Tạo thư mục đầu ra nếu cần
    output_dirs = [os.path.dirname(args.index_path), os.path.dirname(args.documents_path)]
    for directory in output_dirs:
        if directory and not os.path.exists(directory):
            os.makedirs(directory, exist_ok=True)
            logger.info(f"Đã tạo thư mục: {directory}")

    # Chuyển đổi extensions thành list
    extensions = [ext.strip() for ext in args.extensions.split(",")]

    # Hiển thị thông tin cấu hình
    logger.info("=== Cấu hình xây dựng Vector Database ===")
    logger.info(f"Thư mục tài liệu: {args.directory}")
    logger.info(f"Metadata file: {args.metadata if args.metadata else 'Không có'}")
    logger.info(f"FAISS index path: {args.index_path}")
    logger.info(f"Documents path: {args.documents_path}")
    logger.info(f"Model: {args.model_name}")
    logger.info(f"Kích thước chunk: {args.chunk_size}")
    logger.info(f"Độ chồng lấp: {args.chunk_overlap}")
    logger.info(f"Định dạng file: {extensions}")
    logger.info(f"Loại index: {args.index_type}")
    logger.info("=====================================")

    # Gọi hàm tạo vector database
    try:
        create_vector_database(
            directory=args.directory,
            metadata_path=args.metadata,
            index_path=args.index_path,
            documents_path=args.documents_path,
            model_name=args.model_name,
            chunk_size=args.chunk_size,
            chunk_overlap=args.chunk_overlap,
            extensions=extensions,
            index_type=args.index_type
        )
        logger.info("Quá trình xây dựng Vector Database hoàn tất!")
    except KeyboardInterrupt:
        logger.warning("Quá trình bị dừng bởi người dùng.")
        sys.exit(130)
    except Exception as e:
        logger.error(f"Lỗi khi xây dựng Vector Database: {e}")
        import traceback

        traceback.print_exc()
        sys.exit(1)
